#!/usr/bin/env python3
import json, os, sys
sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from renderers_docx import (add_para, add_image, add_hand_table, add_bidding_table,
                             add_generic_table, add_section_title_bar, add_blocks,
                             add_article_header, add_heading_bar, NAVY, BURGUNDY, GOLD, INK_SOFT, WHITE)

BUILD = os.path.dirname(os.path.abspath(__file__))
CONTENT = os.path.join(BUILD, "content")
MEDIA = os.path.join(BUILD, "media")

def load(name):
    with open(os.path.join(CONTENT, f"{name}.json"), encoding="utf-8") as f:
        return json.load(f)

main_blocks = load("main")
gonggao_blocks = load("gonggao")
spingold_blocks = load("spingold_fixed")
xiaocai_blocks = load("xiaocai")
manguan_blocks = load("manguan")
jingque2c_blocks = load("jingque2c")
essays = json.load(open(os.path.join(BUILD, "essays.json"), encoding="utf-8"))

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "PMingLiU"
normal.font.size = Pt(10.5)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
if rFonts is None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
from docx.oxml.ns import qn as _qn
rFonts.set(_qn('w:eastAsia'), 'PMingLiU')

sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)

def page_break():
    doc.add_page_break()

def main_slice(a, b):
    return main_blocks[a:b]

def article(title, author, blocks, media_dir, kicker="", mini_table=False):
    add_article_header(doc, title, author, kicker)
    add_blocks(doc, blocks, media_dir, mini_table=mini_table)

# ============ COVER ============
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("C H U N G   C H I A O   B R I D G E")
r.font.size = Pt(11)
r.font.color.rgb = GOLD
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("中橋季刊")
r.font.size = Pt(54)
r.font.color.rgb = NAVY
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026　秋季號")
r.font.size = Pt(18)
r.font.color.rgb = BURGUNDY

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("中華民國橋藝協會　發行")
r.font.size = Pt(10.5)

for _ in range(2):
    doc.add_paragraph()

highlights = [
    "捷報：第五屆亞洲盃橋藝錦標賽喜獲兩金",
    "橋壇人物專訪：亞洲橋王黃光輝．理事長沈乃正",
    "萬事起頭難——談夢家的第一磴出牌／沈乃正",
    "世青賽專題：合肥世界青年橋牌錦標賽　選手心得與教練講評",
    "橋藝教室：精確2C系統改良、低花黑木氏8S9C法、Bridge Coups",
]
for h in highlights:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.size = Pt(10.5)

for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("中華民國115年10月3日出版　｜　發行人｜沈乃正　主編｜楊欣龍")
r.font.size = Pt(9)
r.font.color.rgb = INK_SOFT

page_break()

# ============ TOC ============
def toc_section(title, items):
    add_section_title_bar(doc, title)
    for t, a in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(t)
        r1.font.size = Pt(10.5)
        r2 = p.add_run("　"+a)
        r2.font.size = Pt(9)
        r2.font.color.rgb = INK_SOFT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("目　　錄")
r.font.size = Pt(22)
r.bold = True
r.font.color.rgb = NAVY
doc.add_paragraph().paragraph_format.space_after = Pt(6)

toc_section("焦點快訊", [
    ("捷報！第五屆亞洲盃橋藝錦標賽喜獲兩金", "本刊報導"),
    ("捷報！第二屆亞洲雙人錦標賽傳佳績", "本刊報導"),
    ("2026亞洲雙人錦標賽介紹", "Doris"),
    ("香港城市橋牌錦標賽", "本刊收錄"),
    ("2026嘉興台灣青年橋藝交流活動", "楊欣龍"),
])
toc_section("橋海漫遊", [
    ("弔顏永南橋友", "沈乃正"),
    ("數風流人物·橋壇人物專訪——黃光輝", "許君薇"),
    ("數風流人物·橋壇人物專訪——沈乃正", "許君薇"),
    ("萬事起頭難——談夢家的第一磴出牌", "沈乃正"),
    ("中曾盃滿貫大戰", "楊欣龍"),
    ("驚技之國之櫻雪橋記（上）", "潘正全"),
    ("續胡說橋牌", "胡希真"),
    ("小材大用", "沈乃正"),
    ("滿貫滿天飛——橋協盃混合組記趣", "洪乙安"),
    ("2026 Summer NABC－Spingold史賓果淘汰賽", "渡鴉"),
    ("經典重刊：第十六屆大專盃橋賽", "徐嘉華．傅依俊"),
])
toc_section("世青賽專題．合肥紀行", [("十六位選手與家長的合肥心得", "本刊採訪整理")])
toc_section("橋藝教室", [
    ("【專欄】Bridge Master解析　水平1：18～23", "陳飛"),
    ("特約介紹——轉換叫後無王開叫者的大配合", "楊欣龍"),
    ("低花黑木氏改良：8S9C法", "任乘風"),
    ("Bridge Coups－4", "黃光輝"),
    ("精確2C後續系統性改進", "上海橋友投稿"),
])
toc_section("橋訊公告", [
    ("中華民國橋藝協會年度重要活動時程", "秘書處"),
    ("各地例行橋藝活動聯絡資訊", "秘書處"),
    ("2026世界跨國邀請賽／世界青年跨國錦標賽成績報告", "秘書處．本刊彙整"),
    ("中華橋協重要比賽成績公告", "秘書處"),
    ("中華民國橋藝協會志工招募", "秘書處"),
    ("徵稿啟事暨版權資料", "本刊敬啟"),
])

page_break()

# ============ SECTION 1: 焦點快訊 ============
add_heading_bar(doc, "第一章　焦點快訊　LATEST NEWS & HIGHLIGHTS", "7C1830", WHITE, Pt(16))
doc.add_paragraph()

article("捷報！第五屆亞洲盃橋藝錦標賽喜獲兩金", "本刊報導", main_slice(3, 6), MEDIA + "/main", kicker="焦點快訊")
article("捷報！第二屆亞洲雙人錦標賽傳佳績", "本刊報導", main_slice(7, 8), MEDIA + "/main", kicker="焦點快訊")
page_break()
article("2026亞洲雙人錦標賽介紹", "Doris", main_slice(9, 15), MEDIA + "/main", kicker="焦點快訊")
page_break()
article("香港城市橋牌錦標賽", "本刊收錄", main_slice(16, 19), MEDIA + "/main", kicker="焦點快訊")
page_break()
article("2026嘉興台灣青年橋藝交流活動", "楊欣龍", main_slice(20, 53), MEDIA + "/main", kicker="焦點快訊")
page_break()

# ============ SECTION 2: 橋海漫遊 ============
add_heading_bar(doc, "第二章　橋海漫遊　STORIES FROM THE BRIDGE WORLD", "7C1830", WHITE, Pt(16))
doc.add_paragraph()

article("弔顏永南橋友", "沈乃正", main_slice(55, 77), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("數風流人物·橋壇人物專訪——黃光輝", "許君薇", main_slice(78, 98), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("數風流人物·橋壇人物專訪——沈乃正", "許君薇", main_slice(99, 130), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("萬事起頭難——談夢家的第一磴出牌", "沈乃正", main_slice(131, 203), MEDIA + "/main", kicker="橋海漫遊", mini_table=True)
page_break()
article("中曾盃滿貫大戰", "楊欣龍", main_slice(204, 239), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("驚技之國之櫻雪橋記（上）", "潘正全", main_slice(240, 294), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("續胡說橋牌", "胡希真", main_slice(295, 357), MEDIA + "/main", kicker="橋海漫遊")
page_break()
article("小材大用", "沈乃正", xiaocai_blocks, MEDIA + "/xiaocai", kicker="橋海漫遊")
page_break()
article("滿貫滿天飛——橋協盃混合組記趣", "洪乙安", manguan_blocks[1:], MEDIA + "/manguan", kicker="橋海漫遊")
page_break()
article("2026 Summer NABC－Spingold史賓果淘汰賽", "渡鴉", spingold_blocks[1:], MEDIA + "/spingold", kicker="橋海漫遊")
page_break()

# archival reprint
add_article_header(doc, "第十六屆大專盃橋賽", "徐嘉華．傅依俊", "橋海漫遊．經典重刊")
note = doc.add_paragraph()
r = note.add_run("經典重刊．典藏自民國橋壇文獻，原載於《第十六屆大專盃橋藝錦標賽》紀念特刊，特此重刊以饗讀者。")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x5c, 0x44, 0x15)
r.italic = True
archival_paras = [
    "第十六屆大專盃橋藝錦標賽，輪由東海大學橋藝社主辦，東海大學是在台中市的郊區，地處偏僻，交通不便，增加與賽橋友不少困擾，加之我們這群從無舉辦橋賽的經驗，籌劃難免欠週，服務不到之處，尚請各校橋友原諒。",
    "本屆橋賽，參加公開組者三十八隊，女子組十一隊，近四百橋友濟濟一堂，場面浩大，在中部地區，可謂少見，公開組分初賽和決賽，女子組單循環比賽，另舉行論對賽，亦分初賽與決賽，其戰況如下：",
    "壹、公開組",
    "三十八隊分四組，計黑桃（S）紅心（H）方塊（D）及梅花（C）。S、H各十隊，D、C各九隊，每組單循環，各取前兩名晉入決賽。",
    "S組：東海東與中興商賽至最後一牌，前者始以剛好的比數（5:3）進入決賽。同組世新一路順風直入。",
    "H組：清華同樣情形，最後一場對中正險勝，以0.5勝分氣走東海風。同組高醫以分組冠軍入。",
    "D組：中興法、T.M.C以一路領先進入決賽。",
    "C組：B.T.U.T險遭滑鐵盧，情況如下：若政大、淡江最後一場均拿足7分以上，則台大回生乏術，結果，政大對中山未拿足，飲恨北歸、淡江、台大晉級。",
    "一決　賽一",
    "先是東海東、中興法最有希望，二隊碰頭，中興法果然犀利而東海東牌運欠佳，加之調配失當，以8:-1飲恨，大失奪標希望。至最後二場有希望問鼎的有中興法、B.T.U.T、東海東、清華；結果清華力克中興法，B.T.U.T.拿8分，東海東雖以4:4平勝高醫，已與冠軍絕緣，最後一場，決定冠、亞、季、殿軍；東海東拿8分，坐三望二，清華-2:8輸給T.M.C.敬坐殿軍，B.T.U.T對中興法，如B.T.U.T贏足則冠軍，可惜以2:6敗北，只得退居季軍。",
    "貳、女子組",
    "女子組最初幾場，尚不能看出誰高誰下，直到最後二場始明朗化：B.T.U.B穩第四，中興慧坐三望二，B.T.U.T與政大爭后，苦戰三小時餘，B.T.U.T技高一籌，政大落至季軍。",
    "參、論對賽",
    "論對賽初賽，採密契爾制，分A、B兩組每組二十一桌，各取南北、東西前四名晉入決賽，決賽採浩威爾制，比賽結果如下：",
    "1. 冠軍：李龍秋、潘國昭（交通大學）　　2. 亞軍：陳全勇、葉建朝（政治大學）",
    "3. 季軍：范熙揚、黃志中（交通大學）　　4. 殿軍：鄭長榮、施慶賢（輔仁大學）",
    "5. 第五名：彭火龍、楊錢棟（文化學院）",
    "本屆賽事由救國團台中地區大專學生社團服務中心協辦，於四月二日至六日假東海大學體育館舉行，宗旨在切磋橋藝、增進友誼；比賽以中國橋藝協會一九七二年出版之複式橋規為比賽規則，公開組取六名、女子組取四名，各頒優勝獎及紀念品，冠軍主盃須連續獲得三屆冠軍，始得永久保留。",
    "（原文作者：徐嘉華、傅依俊；隊職員名單、決賽對戰紀錄等原始文獻，因年代久遠、原稿部分手跡漫漶，本刊謹擇要重刊上述文字紀錄，隊伍名單詳見原刊影本存檔，特此說明並向兩位原作者致謝。）",
]
for para in archival_paras:
    if para in ("壹、公開組", "貳、女子組", "參、論對賽", "一決　賽一"):
        add_para(doc, para, color=NAVY, bold=True, size=Pt(11))
    else:
        add_para(doc, para)
page_break()

# ============ SECTION 3: 世青賽專題 ============
add_heading_bar(doc, "第三章　世青賽專題　2026 WORLD YOUTH TEAMS · HEFEI", "7C1830", WHITE, Pt(16))
doc.add_paragraph()
add_article_header(doc, "合肥紀行——十六位選手與家長的心得", "本刊編輯部　講評｜帶隊教練", "世青賽專題")
add_para(doc, "2026年8月，逐夢隊一行十餘位青少年選手，跟隨教練遠赴中國合肥，參加世界青年橋牌錦標賽。這是許多孩子生平第一次站上國際賽場，緊張與期待交織。賽後，本刊邀請選手與隨隊家長寫下心得，並特邀資深教練逐篇講評——不只是牌技的檢討，更是心理素質、團隊默契與挫折復原力的成長紀錄。")
doc.add_paragraph()

for e in essays:
    add_heading_bar(doc, f'{e["name"]}　｜　{e["role"]}', "10203F", WHITE, Pt(12))
    lbl = doc.add_paragraph()
    r = lbl.add_run("賽前想像")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GOLD
    add_para(doc, e["before"], color=INK_SOFT, italic=True, size=Pt(9.5))
    lbl2 = doc.add_paragraph()
    r2 = lbl2.add_run("參賽心得")
    r2.bold = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GOLD
    add_para(doc, e["after"])
    lbl3 = doc.add_paragraph()
    r3 = lbl3.add_run("教練講評")
    r3.bold = True
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = BURGUNDY
    add_para(doc, e["comment"], size=Pt(9.8))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

page_break()

# ============ SECTION 4: 橋藝教室 ============
add_heading_bar(doc, "第四章　橋藝教室　THE BRIDGE CLASSROOM", "7C1830", WHITE, Pt(16))
doc.add_paragraph()

article("【專欄】Bridge Master解析　水平1：18～23", "陳飛", main_slice(359, 413), MEDIA + "/main", kicker="橋藝教室")
page_break()
article("特約介紹——轉換叫後無王開叫者的大配合", "楊欣龍", main_slice(414, 462), MEDIA + "/main", kicker="橋藝教室")
page_break()
article("低花黑木氏改良：8S9C法", "任乘風", main_slice(463, 469), MEDIA + "/main", kicker="橋藝教室")
page_break()
article("Bridge Coups－4", "黃光輝", main_slice(470, 521), MEDIA + "/main", kicker="橋藝教室")
page_break()
article("精確2C後續系統性改進", "上海橋友投稿", jingque2c_blocks[1:], MEDIA + "/jingque2c", kicker="橋藝教室")
page_break()

# ============ SECTION 5: 橋訊公告 ============
add_heading_bar(doc, "第五章　橋訊公告　ANNOUNCEMENTS & NOTICES", "7C1830", WHITE, Pt(16))
doc.add_paragraph()

g = gonggao_blocks
article("中華民國橋藝協會年度重要活動時程", "秘書處", g[1:3], MEDIA + "/gonggao", kicker="橋訊公告")
page_break()
article("各地例行橋藝活動聯絡資訊", "秘書處", g[3:6], MEDIA + "/gonggao", kicker="橋訊公告")
page_break()
article("2026世界跨國邀請賽／世青賽成績報告", "秘書處．本刊彙整", g[6:10], MEDIA + "/gonggao", kicker="橋訊公告")
page_break()
article("中華橋協重要比賽成績公告", "秘書處", g[10:11] + g[12:14], MEDIA + "/gonggao", kicker="橋訊公告")
page_break()
article("中華民國橋藝協會志工招募暨徵稿啟事", "秘書處", g[15:30], MEDIA + "/gonggao", kicker="橋訊公告")
page_break()

# ============ COLOPHON ============
add_heading_bar(doc, "版權頁　COLOPHON", "10203F", WHITE, Pt(14))
doc.add_paragraph()
colophon_rows = [
    ("發行人", "沈乃正"), ("主編", "楊欣龍　shon.yang@gmail.com"),
    ("出版者", "中華民國橋藝協會　ctcba886@gmail.com"),
    ("地址", "臺北市大安區忠孝東路三段217巷7弄7號B1"),
    ("電話", "(02)2772-4510、(02)2772-4583"), ("傳真", "(02)2772-4493"),
    ("橋協網站", "https://www.ctcba.org.tw/"), ("文稿校對", "陸怡如"),
    ("出版日期", "中華民國115年10月3日"),
]
add_generic_table(doc, [[k, v] for k, v in colophon_rows])
add_para(doc, "本刊採電子版發行，歡迎海內外橋友踴躍投稿以饗讀者　｜　版權所有．轉載請註明出處", color=INK_SOFT, size=Pt(9))

out_path = os.path.expanduser("~/Desktop/AI同事/中橋季刊_2026秋季號.docx")
doc.save(out_path)
print("Saved", out_path)
