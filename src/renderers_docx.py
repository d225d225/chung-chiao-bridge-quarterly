#!/usr/bin/env python3
"""python-docx helpers mirroring renderers.py, for the native Word/Pages build."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

NAVY = RGBColor(0x10, 0x20, 0x3f)
BURGUNDY = RGBColor(0x7c, 0x18, 0x30)
GOLD = RGBColor(0xb6, 0x89, 0x3f)
INK = RGBColor(0x1c, 0x1c, 0x1c)
INK_SOFT = RGBColor(0x4a, 0x4a, 0x4a)
CREAM = "FAF6EE"
WHITE = RGBColor(0xff, 0xff, 0xff)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)

def add_suit_runs(paragraph, text, base_color=None, size=None, bold=False, italic=False):
    for ch in text:
        run = paragraph.add_run(ch)
        run.font.size = size or Pt(10.5)
        run.bold = bold
        run.italic = italic
        if ch in "♥♦":
            run.font.color.rgb = BURGUNDY
        elif ch in "♠♣":
            run.font.color.rgb = NAVY
        elif base_color is not None:
            run.font.color.rgb = base_color

def add_para(doc, text, color=None, size=Pt(10.5), bold=False, italic=False, align=None, space_after=Pt(6)):
    for seg in (text or "").split("\n"):
        seg = seg.strip()
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        if seg:
            add_suit_runs(p, seg, base_color=color, size=size, bold=bold, italic=italic)
        else:
            p.add_run("")
    return None

def add_heading_bar(doc, text, bg_hex, fg_color=WHITE, size=Pt(12)):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = fg_color
    return table

def add_image(doc, path, max_width_in=6.2):
    if not os.path.exists(path):
        return
    try:
        with Image.open(path) as im:
            w, h = im.size
        ratio = h / w if w else 1
        width = Inches(max_width_in)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    except Exception:
        pass

def add_hand_table(doc, data):
    table = doc.add_table(rows=3, cols=3)
    table.autofit = True
    for r in range(3):
        for c in range(3):
            set_cell_border_none(table.cell(r, c))
            table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def fill(cell, lines, align_center=True):
        p = cell.paragraphs[0]
        if align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for i, line in enumerate(lines):
            if i > 0:
                p = cell.add_paragraph()
                if align_center:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
            add_suit_runs(p, line, size=Pt(9.5))

    label = " ／ ".join(data.get("label", []))
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(label)
        r.font.size = Pt(8.5)
        r.font.color.rgb = BURGUNDY
        cap.paragraph_format.space_after = Pt(2)

    fill(table.cell(0, 1), data.get("north", []))
    fill(table.cell(1, 0), data.get("west", []))
    fill(table.cell(1, 2), data.get("east", []))
    fill(table.cell(2, 1), data.get("south", []))
    center = table.cell(1, 1)
    cp = center.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("N\nW  E\nS")
    r.font.size = Pt(7)
    r.font.color.rgb = INK_SOFT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_bidding_table(doc, data):
    header = data.get("header", ["W", "N", "E", "S"])
    rows = data.get("rows", [])
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    for c, h in enumerate(header):
        cell = table.cell(0, c)
        shade_cell(cell, "10203F")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for ri, r in enumerate(rows, start=1):
        for c, h in enumerate(header):
            cell = table.cell(ri, c)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_suit_runs(p, r.get(h, ""), size=Pt(9))
    for fn in data.get("footnotes", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        add_suit_runs(p, fn, base_color=INK_SOFT, size=Pt(9), italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_generic_table(doc, grid):
    if not grid:
        return
    ncols = max((len(r) for r in grid), default=0)
    if ncols == 2:
        for r in grid:
            if len(r) == 1:
                p = doc.add_paragraph()
                run = p.add_run(r[0])
                run.bold = True
                run.font.color.rgb = WHITE
                shade_cell_p = doc.add_table(rows=1, cols=1).rows[0].cells[0]
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run((r[0] or "") + "：")
                run.bold = True
                run.font.color.rgb = BURGUNDY
                run.font.size = Pt(9.5)
                val = (r[1] if len(r) > 1 else "").replace("\n", "  ")
                run2 = p.add_run(val)
                run2.font.size = Pt(9.5)
        return
    table = doc.add_table(rows=len(grid), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(grid):
        cells = (row + [""] * ncols)[:ncols]
        for ci, val in enumerate(cells):
            cell = table.cell(ri, ci)
            if ri == 0:
                shade_cell(cell, "7C1830")
            p = cell.paragraphs[0]
            for i, line in enumerate(val.split("\n")):
                if i > 0:
                    p = cell.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(8.5)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_section_title_bar(doc, text):
    add_heading_bar(doc, text, "7C1830", WHITE, Pt(13))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_block(doc, b, media_dir, mini_table=False):
    t = b["type"]
    if t == "para":
        add_para(doc, b.get("text", ""))
        for im in b.get("images", []) or []:
            add_image(doc, os.path.join(media_dir, im))
    elif t == "heading":
        add_para(doc, b["text"], color=BURGUNDY, size=Pt(12), bold=True, space_after=Pt(4))
        for im in b.get("images", []) or []:
            add_image(doc, os.path.join(media_dir, im))
    elif t == "title":
        add_para(doc, b["text"], color=NAVY, size=Pt(10.5), bold=True, space_after=Pt(4))
    elif t == "hand":
        add_hand_table(doc, b["data"])
    elif t == "bidding":
        add_bidding_table(doc, b["data"])
    elif t == "table":
        add_generic_table(doc, b["data"])
    elif t == "note":
        add_para(doc, b.get("text", ""), color=INK_SOFT, italic=True, size=Pt(9.5))

def add_blocks(doc, blocks, media_dir, mini_table=False):
    for b in blocks:
        add_block(doc, b, media_dir, mini_table=mini_table)

def add_article_header(doc, title, author, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        r = p.add_run(kicker)
        r.font.size = Pt(8.5)
        r.font.color.rgb = GOLD
        r.bold = True
        p.paragraph_format.space_after = Pt(2)
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(2)
    b = doc.add_paragraph()
    r = b.add_run(f"文｜{author}")
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK_SOFT
    b.paragraph_format.space_after = Pt(8)
